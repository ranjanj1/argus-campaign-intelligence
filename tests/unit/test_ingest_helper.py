"""Unit tests for IngestHelper and TableChunker."""
from __future__ import annotations

import csv
import io
from pathlib import Path
from unittest.mock import MagicMock, patch

import pytest

from argus.components.ingest.ingest_helper import IngestHelper, ParsedPage
from argus.components.ingest.table_chunker import TableChunker


# ── IngestHelper ─────────────────────────────────────────────────────────────

class TestIngestHelperPDF:
    def test_parse_pdf_returns_pages(self, tmp_path):
        helper = IngestHelper()
        # Patch _parse_pdf so no real file I/O is needed
        with patch.object(helper, "_parse_pdf") as mock:
            mock.return_value = [ParsedPage(text="page one content", page=1)]
            result = helper.parse(tmp_path / "report.pdf")
        assert len(result) == 1
        assert result[0].page == 1
        assert "page one" in result[0].text

    def test_parse_pdf_bytes(self):
        helper = IngestHelper()
        with patch.object(helper, "_parse_pdf_stream") as mock:
            mock.return_value = [ParsedPage(text="hello", page=1)]
            result = helper.parse_bytes(b"%PDF-1.4", "test.pdf")
        assert result[0].text == "hello"


class TestIngestHelperDOCX:
    def test_parse_docx(self, tmp_path):
        from docx import Document
        doc = Document()
        doc.add_paragraph("First paragraph")
        doc.add_paragraph("Second paragraph")
        path = tmp_path / "brief.docx"
        doc.save(str(path))

        helper = IngestHelper()
        pages = helper.parse(path)
        assert len(pages) == 1
        assert "First paragraph" in pages[0].text
        assert "Second paragraph" in pages[0].text

    def test_parse_docx_with_table(self, tmp_path):
        from docx import Document
        doc = Document()
        table = doc.add_table(rows=2, cols=2)
        table.rows[0].cells[0].text = "Campaign"
        table.rows[0].cells[1].text = "ROAS"
        table.rows[1].cells[0].text = "Summer Push"
        table.rows[1].cells[1].text = "4.2"
        path = tmp_path / "table.docx"
        doc.save(str(path))

        helper = IngestHelper()
        pages = helper.parse(path)
        assert "Campaign" in pages[0].text
        assert "4.2" in pages[0].text


class TestIngestHelperCSV:
    def test_parse_csv(self, tmp_path):
        path = tmp_path / "data.csv"
        path.write_text("name,value\nfoo,1\nbar,2\n")
        helper = IngestHelper()
        pages = helper.parse(path)
        assert len(pages) == 1
        assert "foo" in pages[0].text

    def test_parse_csv_bytes(self):
        helper = IngestHelper()
        data = b"col1,col2\nA,B\n"
        pages = helper.parse_bytes(data, "data.csv")
        assert "col1" in pages[0].text


class TestIngestHelperXLSX:
    def test_parse_xlsx(self, tmp_path):
        import openpyxl
        wb = openpyxl.Workbook()
        ws = wb.active
        ws.title = "Sheet1"
        ws.append(["campaign", "spend", "roas"])
        ws.append(["Alpha", 10000, 3.5])
        path = tmp_path / "budget.xlsx"
        wb.save(str(path))

        helper = IngestHelper()
        pages = helper.parse(path)
        assert len(pages) == 1
        assert "campaign" in pages[0].text.lower()
        assert pages[0].extra["sheet"] == "Sheet1"


class TestIngestHelperTXT:
    def test_parse_txt(self, tmp_path):
        path = tmp_path / "notes.txt"
        path.write_text("Some strategy notes here.")
        helper = IngestHelper()
        pages = helper.parse(path)
        assert pages[0].text == "Some strategy notes here."

    def test_unsupported_extension_raises(self, tmp_path):
        helper = IngestHelper()
        with pytest.raises(ValueError, match="Unsupported"):
            helper.parse(tmp_path / "data.parquet")


# ── TableChunker ──────────────────────────────────────────────────────────────

class TestTableChunker:
    def test_basic_chunking(self):
        csv_text = _make_csv(header=["a", "b"], rows=25)
        chunker = TableChunker()
        chunks = chunker.chunk_csv(csv_text, rows_per_chunk=10)
        assert len(chunks) == 3   # 10 + 10 + 5
        assert chunks[0].row_start == 1
        assert chunks[0].row_end == 10
        assert chunks[1].row_start == 11
        assert chunks[2].row_end == 25

    def test_header_in_every_chunk(self):
        csv_text = _make_csv(header=["col1", "col2"], rows=15)
        chunks = TableChunker().chunk_csv(csv_text, rows_per_chunk=5)
        for chunk in chunks:
            first_line = chunk.text.split("\n")[0]
            assert "col1" in first_line
            assert "col2" in first_line

    def test_empty_csv(self):
        chunks = TableChunker().chunk_csv("", rows_per_chunk=10)
        assert chunks == []

    def test_header_only_csv(self):
        chunks = TableChunker().chunk_csv("col1,col2\n", rows_per_chunk=10)
        assert chunks == []

    def test_single_row(self):
        chunks = TableChunker().chunk_csv("a,b\n1,2\n", rows_per_chunk=10)
        assert len(chunks) == 1
        assert chunks[0].row_start == 1
        assert chunks[0].row_end == 1

    def test_sheet_name_propagated(self):
        csv_text = _make_csv(header=["x"], rows=5)
        chunks = TableChunker().chunk_xlsx_sheet(
            csv_text, sheet_name="Budget", rows_per_chunk=10
        )
        assert chunks[0].extra["sheet"] == "Budget"


# ── helpers ───────────────────────────────────────────────────────────────────

def _make_csv(header: list[str], rows: int) -> str:
    buf = io.StringIO()
    writer = csv.writer(buf)
    writer.writerow(header)
    for i in range(rows):
        writer.writerow([f"val{j}_{i}" for j in range(len(header))])
    return buf.getvalue()
