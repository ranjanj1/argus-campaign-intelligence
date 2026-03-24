from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from jinja2 import Environment, FileSystemLoader

from scripts.models.client_context import ClientProfile

TEMPLATES_DIR = Path(__file__).parent.parent / "templates"

BRAND_BLUE = RGBColor(0x1F, 0x4E, 0x79)
BRAND_LIGHT = RGBColor(0x2E, 0x75, 0xB6)


def _build_brief_context(profile: ClientProfile) -> dict:
    config = profile.vertical_config
    segments_summary = [
        f"{s.segment_name} ({s.age_range}, {s.platform})"
        for s in profile.segments[:6]
    ]
    return {
        "client_id": profile.client_id,
        "company_name": profile.company_name,
        "industry": profile.industry,
        "primary_channels": config.primary_channels,
        "monthly_budget_min": config.monthly_budget_range[0],
        "monthly_budget_max": config.monthly_budget_range[1],
        "segments_summary": segments_summary,
        "business_objectives": config.business_objectives,
        "competitors": config.competitors,
        "tone_keywords": config.tone_keywords,
    }


def _add_heading(doc: Document, text: str, level: int = 1) -> None:
    para = doc.add_heading(text, level=level)
    for run in para.runs:
        run.font.color.rgb = BRAND_BLUE if level == 1 else BRAND_LIGHT
        run.font.size = Pt(16 if level == 1 else 13)


def _add_body(doc: Document, text: str) -> None:
    for para_text in text.strip().split("\n\n"):
        para_text = para_text.strip()
        if para_text:
            p = doc.add_paragraph(para_text)
            for run in p.runs:
                run.font.size = Pt(10.5)
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


def _add_bullet(doc: Document, text: str) -> None:
    p = doc.add_paragraph(text.lstrip("- "), style="List Bullet")
    for run in p.runs:
        run.font.size = Pt(10.5)


def _add_kpi_table(doc: Document, kpi_rows: list[str]) -> None:
    """kpi_rows: each string is 'KPI | Target | Method'"""
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for cell, label in zip(hdr, ["KPI", "Target", "Measurement Method"]):
        cell.text = label
        for para in cell.paragraphs:
            for run in para.runs:
                run.font.bold = True
                run.font.size = Pt(10)
    for row_str in kpi_rows:
        parts = [p.strip() for p in row_str.split("|")]
        if len(parts) >= 3:
            row = table.add_row().cells
            for cell, val in zip(row, parts[:3]):
                cell.text = val
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.font.size = Pt(10)


def _render_docx(sections: dict, context: dict, output_path: Path) -> None:
    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    # Cover
    title = doc.add_heading("Campaign Strategy Brief", 0)
    for run in title.runs:
        run.font.color.rgb = BRAND_BLUE
        run.font.size = Pt(22)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    sub = doc.add_paragraph(f"{context['company_name']} | {context['industry'].title()} Vertical")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in sub.runs:
        run.font.color.rgb = BRAND_LIGHT
        run.font.size = Pt(13)

    doc.add_paragraph("Prepared by 2060 Digital").alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()

    _add_heading(doc, "1. Client Background")
    _add_body(doc, sections.get("client_background", ""))

    _add_heading(doc, "2. Campaign Objectives")
    for obj in sections.get("campaign_objectives", []):
        _add_bullet(doc, obj)

    _add_heading(doc, "3. Target Audience")
    _add_body(doc, sections.get("target_audience", ""))

    _add_heading(doc, "4. Channel Strategy")
    _add_body(doc, sections.get("channel_strategy", ""))

    doc.add_page_break()

    _add_heading(doc, "5. Key Performance Indicators")
    kpi_rows = sections.get("kpi_table", [])
    if kpi_rows:
        _add_kpi_table(doc, kpi_rows)
    doc.add_paragraph()

    _add_heading(doc, "6. Budget Rationale")
    _add_body(doc, sections.get("budget_rationale", ""))

    _add_heading(doc, "7. Risk Mitigations")
    for risk in sections.get("risk_mitigations", []):
        _add_bullet(doc, risk)

    if sections.get("success_timeline"):
        _add_heading(doc, "8. Success Timeline")
        _add_body(doc, sections["success_timeline"])

    doc.save(str(output_path))


def generate_strategy_brief_docx(profile: ClientProfile, output_path: Path) -> None:
    """Render DOCX using Jinja2 templates — no LLM required."""
    env = Environment(loader=FileSystemLoader(str(TEMPLATES_DIR)), autoescape=False)
    context = _build_brief_context(profile)
    template = env.get_template("strategy_brief.j2")
    rendered = template.render(**context)

    # Parse sections
    sections: dict[str, str | list] = {}
    current_key = None
    buf: list[str] = []
    for line in rendered.splitlines():
        if line.startswith("###SECTION:"):
            if current_key:
                if isinstance(sections.get(current_key), list):
                    pass
                else:
                    sections[current_key] = "\n".join(buf).strip()
                buf = []
            current_key = line.replace("###SECTION:", "").strip()
        elif line.startswith("###LIST:"):
            if current_key and buf:
                sections[current_key] = "\n".join(buf).strip()
                buf = []
            current_key = line.replace("###LIST:", "").strip()
            sections[current_key] = []
        elif current_key is not None:
            existing = sections.get(current_key)
            if isinstance(existing, list):
                if line.strip():
                    existing.append(line.strip())
            else:
                buf.append(line)
    if current_key and buf and not isinstance(sections.get(current_key), list):
        sections[current_key] = "\n".join(buf).strip()

    _render_docx(sections, context, output_path)
