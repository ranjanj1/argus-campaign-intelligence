from __future__ import annotations

"""
IngestHelper — file parsing and text extraction.

Supports:
  .pdf   → pypdf text extraction (page-by-page)
  .docx  → python-docx paragraph extraction
  .csv   → raw text (rows will be chunked by TableChunker separately)
  .xlsx  → sheet → CSV-like text (also handled by TableChunker)
  .txt   → plain read

Returns a list of ParsedPage objects: one per PDF page or one for the
whole document for non-PDF formats.
"""

import csv
import io
import logging
from dataclasses import dataclass, field
from pathlib import Path
from typing import BinaryIO

logger = logging.getLogger(__name__)

_MAX_PDF_PAGES = 500   # safety cap


@dataclass
class ParsedPage:
    """
    A single logical unit of text extracted from a file.

    For PDFs: one entry per page.
    For DOCX/CSV/XLSX/TXT: one entry for the whole file.
    """
    text: str
    page: int = 0                         # 1-based for PDFs; 0 = not paged
    extra: dict = field(default_factory=dict)  # format-specific metadata


class IngestHelper:
    """
    Stateless file parser — no DI needed.  Instantiate once; call parse().

    Usage:
        helper = IngestHelper()
        pages = helper.parse(Path("q3_report.pdf"))
    """

    def parse(self, path: Path) -> list[ParsedPage]:
        """
        Parse a file and return a list of ParsedPage objects.
        Raises ValueError for unsupported file types.
        """
        suffix = path.suffix.lower()
        if suffix == ".pdf":
            return self._parse_pdf(path)
        if suffix == ".docx":
            return self._parse_docx(path)
        if suffix == ".csv":
            return self._parse_csv(path)
        if suffix in (".xlsx", ".xls"):
            return self._parse_xlsx(path)
        if suffix == ".txt":
            return self._parse_txt(path)
        raise ValueError(f"Unsupported file type: {suffix}")

    def parse_bytes(self, data: bytes, filename: str) -> list[ParsedPage]:
        """Parse from raw bytes (used by the ingest router upload handler)."""
        suffix = Path(filename).suffix.lower()
        buf = io.BytesIO(data)
        if suffix == ".pdf":
            return self._parse_pdf_stream(buf)
        if suffix == ".docx":
            return self._parse_docx_stream(buf)
        if suffix == ".csv":
            text = data.decode("utf-8", errors="replace")
            return [ParsedPage(text=text)]
        if suffix in (".xlsx", ".xls"):
            return self._parse_xlsx_stream(buf)
        if suffix == ".txt":
            return [ParsedPage(text=data.decode("utf-8", errors="replace"))]
        raise ValueError(f"Unsupported file type: {suffix}")

    # ── PDF ───────────────────────────────────────────────────────────────────

    def _parse_pdf(self, path: Path) -> list[ParsedPage]:
        with path.open("rb") as f:
            return self._parse_pdf_stream(f)

    def _parse_pdf_stream(self, stream: BinaryIO) -> list[ParsedPage]:
        from pypdf import PdfReader
        reader = PdfReader(stream)
        pages = []
        for i, page in enumerate(reader.pages[:_MAX_PDF_PAGES], start=1):
            text = page.extract_text() or ""
            text = text.strip()
            if text:
                pages.append(ParsedPage(text=text, page=i))
        if not pages:
            logger.warning("PDF produced no extractable text")
        return pages

    # ── DOCX ──────────────────────────────────────────────────────────────────

    def _parse_docx(self, path: Path) -> list[ParsedPage]:
        with path.open("rb") as f:
            return self._parse_docx_stream(f)

    def _parse_docx_stream(self, stream: BinaryIO) -> list[ParsedPage]:
        from docx import Document
        doc = Document(stream)
        paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        # Also extract table cell text
        for table in doc.tables:
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells if c.text.strip()]
                if cells:
                    paragraphs.append(" | ".join(cells))
        text = "\n\n".join(paragraphs)
        return [ParsedPage(text=text)] if text else []

    # ── CSV ───────────────────────────────────────────────────────────────────

    def _parse_csv(self, path: Path) -> list[ParsedPage]:
        text = path.read_text(encoding="utf-8", errors="replace")
        return [ParsedPage(text=text)]

    # ── XLSX ──────────────────────────────────────────────────────────────────

    def _parse_xlsx(self, path: Path) -> list[ParsedPage]:
        with path.open("rb") as f:
            return self._parse_xlsx_stream(f)

    def _parse_xlsx_stream(self, stream: BinaryIO) -> list[ParsedPage]:
        import openpyxl
        wb = openpyxl.load_workbook(stream, read_only=True, data_only=True)
        pages = []
        for sheet in wb.worksheets:
            buf = io.StringIO()
            writer = csv.writer(buf)
            for row in sheet.iter_rows(values_only=True):
                writer.writerow([str(c) if c is not None else "" for c in row])
            text = buf.getvalue().strip()
            if text:
                pages.append(ParsedPage(text=text, extra={"sheet": sheet.title}))
        return pages

    # ── TXT ───────────────────────────────────────────────────────────────────

    def _parse_txt(self, path: Path) -> list[ParsedPage]:
        text = path.read_text(encoding="utf-8", errors="replace").strip()
        return [ParsedPage(text=text)] if text else []
